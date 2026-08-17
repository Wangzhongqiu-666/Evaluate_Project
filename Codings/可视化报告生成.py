"""Generate corrected charts, a portfolio-ready DOCX/PDF, and presentation notes."""

from __future__ import annotations

import json
from datetime import date
from pathlib import Path
from typing import Any

import matplotlib.pyplot as plt
import pandas as pd
import seaborn as sns
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from wordcloud import WordCloud

from evaluation import QUESTION_MARKERS, evaluate_labels
from project_config import DEFAULT_OUTPUT_DIR, ensure_output_dir


BLUE = "2E74B5"
DARK_BLUE = "0B2545"
MID_BLUE = "1F4D78"
MUTED = "5B6573"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"


def _set_run_font(
    run,
    *,
    latin: str = "Calibri",
    east_asia: str = "Microsoft YaHei",
    size: float | None = None,
    bold: bool | None = None,
    color: str | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_dxa))


def _set_table_geometry(table, widths: list[int], indent: int = 120) -> None:
    if sum(widths) != 9360:
        raise ValueError(f"表格列宽必须合计9360 DXA，当前为{sum(widths)}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout", "w:tblCellMar"):
        old = tbl_pr.find(qn(tag))
        if old is not None:
            tbl_pr.remove(old)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), "9360")
    tbl_pr.append(tbl_w)
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_pr.append(tbl_ind)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    margins = OxmlElement("w:tblCellMar")
    for edge, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    tbl_pr.append(margins)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            _set_cell_width(cell, width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _mark_header_repeat(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])
    _set_run_font(run, size=9, color=MUTED)


class ReportGenerator:
    def __init__(
        self,
        output_dir: str | Path = DEFAULT_OUTPUT_DIR,
        label_workbook: str | Path | None = None,
    ) -> None:
        self.output_dir = ensure_output_dir(output_dir)
        self.summary = self._read_json("分析摘要.json")
        self.data = pd.read_csv(
            self.output_dir / "评论分析结果_保研版.csv", encoding="utf-8-sig"
        )
        self.keywords = pd.read_csv(
            self.output_dir / "关键词统计_保研版.csv", encoding="utf-8-sig"
        )
        if label_workbook:
            evaluate_labels(label_workbook, self.output_dir)
        metrics_path = self.output_dir / "模型评估结果.json"
        self.metrics = (
            json.loads(metrics_path.read_text(encoding="utf-8"))
            if metrics_path.exists()
            else None
        )
        self._configure_plot_fonts()

    def _read_json(self, name: str) -> dict[str, Any]:
        return json.loads((self.output_dir / name).read_text(encoding="utf-8"))

    @staticmethod
    def _configure_plot_fonts() -> None:
        plt.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei", "Arial Unicode MS"]
        plt.rcParams["axes.unicode_minus"] = False
        sns.set_theme(style="whitegrid", font="Microsoft YaHei")

    def create_charts(self) -> None:
        labels = ["正面", "中性", "负面"]
        values = [self.summary["positive"], self.summary["neutral"], self.summary["negative"]]
        colors = ["#2A9D8F", "#7C8FA5", "#E76F51"]
        fig, ax = plt.subplots(figsize=(7.2, 4.2))
        bars = ax.barh(labels[::-1], values[::-1], color=colors[::-1], height=0.58)
        ax.set_title("领域模型预测的评论情绪分布", fontsize=14, pad=12)
        ax.set_xlabel("评论数")
        ax.grid(axis="y", visible=False)
        total = self.summary["unique_records"]
        for bar, value in zip(bars, values[::-1]):
            ax.text(
                value + max(values) * 0.02,
                bar.get_y() + bar.get_height() / 2,
                f"{value}（{value / total:.1%}）",
                va="center",
                fontsize=10,
            )
        ax.set_xlim(0, max(values) * 1.28)
        fig.tight_layout()
        fig.savefig(self.output_dir / "情绪分布_保研版.png", dpi=180, bbox_inches="tight")
        plt.close(fig)

        fig, ax = plt.subplots(figsize=(7.2, 4.2))
        sns.histplot(
            self.data["领域模型分数"], bins=15, kde=True, color="#6BAED6", edgecolor="white", ax=ax
        )
        ax.axvline(0.3, color="#E76F51", linestyle="--", linewidth=1.5, label="负面阈值 0.3")
        ax.axvline(0.7, color="#2A9D8F", linestyle="--", linewidth=1.5, label="正面阈值 0.7")
        ax.set_title("领域模型情感分数分布", fontsize=14, pad=12)
        ax.set_xlabel("情感分数（0-1）")
        ax.set_ylabel("评论数")
        ax.legend(frameon=False)
        fig.tight_layout()
        fig.savefig(self.output_dir / "情感分数分布_保研版.png", dpi=180, bbox_inches="tight")
        plt.close(fig)

        top = self.keywords.head(15).sort_values("频数")
        fig, ax = plt.subplots(figsize=(7.2, 5.0))
        ax.barh(top["关键词"], top["频数"], color="#4C78A8")
        ax.set_title("Top 15 高频关键词（已过滤无意义单字）", fontsize=14, pad=12)
        ax.set_xlabel("频数")
        ax.grid(axis="y", visible=False)
        for y, value in enumerate(top["频数"]):
            ax.text(value + 0.3, y, str(int(value)), va="center", fontsize=9)
        fig.tight_layout()
        fig.savefig(self.output_dir / "高频关键词_保研版.png", dpi=180, bbox_inches="tight")
        plt.close(fig)

        frequencies = {
            str(row["关键词"]): int(row["频数"])
            for _, row in self.keywords.head(100).iterrows()
        }
        font_path = Path("C:/Windows/Fonts/msyh.ttc")
        cloud = WordCloud(
            font_path=str(font_path) if font_path.exists() else None,
            width=1200,
            height=650,
            background_color="white",
            colormap="viridis",
            max_words=100,
            random_state=20260811,
        ).generate_from_frequencies(frequencies)
        cloud.to_file(str(self.output_dir / "关键词词云_保研版.png"))

    def _configure_document(self, document: Document) -> None:
        section = document.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

        style = document.styles["Normal"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.10

        for name, size, color, before, after in (
            ("Heading 1", 16, BLUE, 16, 8),
            ("Heading 2", 13, BLUE, 12, 6),
            ("Heading 3", 12, MID_BLUE, 8, 4),
        ):
            heading = document.styles[name]
            heading.font.name = "Calibri"
            heading._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            heading.font.size = Pt(size)
            heading.font.bold = True
            heading.font.color.rgb = RGBColor.from_string(color)
            heading.paragraph_format.space_before = Pt(before)
            heading.paragraph_format.space_after = Pt(after)
            heading.paragraph_format.keep_with_next = True

        for name in ("List Bullet", "List Number"):
            list_style = document.styles[name]
            list_style.font.name = "Calibri"
            list_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            list_style.font.size = Pt(11)
            list_style.paragraph_format.left_indent = Inches(0.5)
            list_style.paragraph_format.first_line_indent = Inches(-0.25)
            list_style.paragraph_format.space_after = Pt(8)
            list_style.paragraph_format.line_spacing = 1.167

        header = section.header.paragraphs[0]
        header.text = "B站教学评论情感分析  |  个人课程项目复盘"
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in header.runs:
            _set_run_font(run, size=9, color=MUTED)
        footer = section.footer.paragraphs[0]
        _add_page_number(footer)

    @staticmethod
    def _add_body(document: Document, text: str, *, bold_prefix: str | None = None) -> None:
        paragraph = document.add_paragraph()
        if bold_prefix and text.startswith(bold_prefix):
            first = paragraph.add_run(bold_prefix)
            _set_run_font(first, bold=True, color=DARK_BLUE)
            rest = paragraph.add_run(text[len(bold_prefix) :])
            _set_run_font(rest)
        else:
            run = paragraph.add_run(text)
            _set_run_font(run)

    @staticmethod
    def _add_callout(document: Document, text: str) -> None:
        table = document.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        _set_table_geometry(table, [9360])
        cell = table.cell(0, 0)
        _set_cell_shading(cell, CALLOUT)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        _set_run_font(run, size=11, bold=True, color=DARK_BLUE)
        document.add_paragraph().paragraph_format.space_after = Pt(2)

    @staticmethod
    def _add_table(document: Document, headers: list[str], rows: list[list[Any]], widths: list[int]):
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        _set_table_geometry(table, widths)
        _mark_header_repeat(table.rows[0])
        for cell, header in zip(table.rows[0].cells, headers):
            _set_cell_shading(cell, LIGHT_GRAY)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(header)
            _set_run_font(run, bold=True, color=DARK_BLUE)
        for values in rows:
            cells = table.add_row().cells
            for index, (cell, value) in enumerate(zip(cells, values)):
                paragraph = cell.paragraphs[0]
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
                )
                run = paragraph.add_run(str(value))
                _set_run_font(run, size=10.5)
        _set_table_geometry(table, widths)
        document.add_paragraph().paragraph_format.space_after = Pt(2)
        return table

    @staticmethod
    def _add_figure(document: Document, path: Path, caption: str, width: float = 5.9) -> None:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(path), width=Inches(width))
        caption_p = document.add_paragraph()
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_p.paragraph_format.space_before = Pt(4)
        caption_p.paragraph_format.space_after = Pt(8)
        run = caption_p.add_run(caption)
        _set_run_font(run, size=9, color=MUTED)

    def create_docx(self) -> Path:
        document = Document()
        self._configure_document(document)
        document.core_properties.title = "B站教学评论情感分析项目报告（保研版）"
        document.core_properties.subject = "中文NLP课程项目复盘"
        document.core_properties.author = "课程项目作者"

        kicker = document.add_paragraph()
        kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
        kicker.paragraph_format.space_before = Pt(28)
        kicker.paragraph_format.space_after = Pt(10)
        run = kicker.add_run("COURSE PROJECT  /  NLP CASE STUDY")
        _set_run_font(run, size=10, bold=True, color=BLUE)

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_after = Pt(8)
        run = title.add_run("B站教学评论情感分析")
        _set_run_font(run, size=28, bold=True, color=DARK_BLUE)

        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.paragraph_format.space_after = Pt(18)
        run = subtitle.add_run("从评论采集、领域情感建模到自动化报告生成")
        _set_run_font(run, size=14, color=MID_BLUE)

        meta = document.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.paragraph_format.space_after = Pt(24)
        run = meta.add_run(f"个人独立完成  |  Python 3.13  |  更新于 {date.today():%Y-%m-%d}")
        _set_run_font(run, size=10, color=MUTED)

        self._add_callout(
            document,
            f"项目形成了“采集 - 清洗 - 分词 - 领域情感分析 - 可视化 - 报告”的完整闭环；"
            f"本案例共分析 {self.summary['raw_records']} 条原始评论，去重后保留 "
            f"{self.summary['unique_records']} 条。",
        )

        document.add_heading("1. 项目目标与系统流程", level=1)
        self._add_body(
            document,
            "目标：面向B站教学视频评论，自动提取用户态度与关注主题，为课程内容优化提供结构化参考。项目定位是单视频案例研究，不将结果外推为通用教学质量结论。",
            bold_prefix="目标：",
        )
        for text in (
            "解析视频标识并按页采集顶层评论，保存昵称、性别与评论正文。",
            "对评论做去重和文本清洗，使用Jieba及教学领域词典进行分词。",
            "加载领域SnowNLP模型，按0.3/0.7阈值划分负面、中性和正面。",
            "统计高频关键词、情绪分布和连续情感分数，并自动生成报告。",
            "构建120条随机测试样本与30条诊断样本，等待独立人工标签后评价模型。",
        ):
            paragraph = document.add_paragraph(style="List Number")
            run = paragraph.add_run(text)
            _set_run_font(run)

        document.add_heading("2. 数据与分析结果", level=1)
        stats_rows = [
            ["原始评论", self.summary["raw_records"], "爬虫输出中的全部记录"],
            ["去重评论", self.summary["unique_records"], f"删除{self.summary['duplicates_removed']}条完全重复评论"],
            ["正面预测", self.summary["positive"], f"{self.summary['positive_share']:.1%}"],
            ["中性预测", self.summary["neutral"], f"{self.summary['neutral_share']:.1%}"],
            ["负面预测", self.summary["negative"], f"{self.summary['negative_share']:.1%}"],
            ["净情感值", f"{self.summary['net_sentiment_points']:+.2f}个百分点", "正面占比减负面占比"],
        ]
        self._add_table(document, ["指标", "结果", "说明"], stats_rows, [2700, 2100, 4560])
        self._add_body(
            document,
            "说明：新版不再使用原报告中经过平方根放大的非标准满意度分数，避免用缺乏依据的单一分数替代原始情绪分布。",
            bold_prefix="说明：",
        )
        self._add_figure(document, self.output_dir / "情绪分布_保研版.png", "图1  去重评论的模型预测分布")
        self._add_figure(document, self.output_dir / "情感分数分布_保研版.png", "图2  连续情感分数及分类阈值")

        document.add_heading("3. 方法设计", level=1)
        method_rows = [
            ["领域词典", "138条教学场景词语", "增强课程、知识点与课堂评价相关分词"],
            ["停用词", "923条", "在循环外一次加载，降低重复I/O"],
            ["训练语料", "正面16,720 / 负面18,976", "训练领域SnowNLP二分类概率模型"],
            ["三分类规则", "<=0.3负面；>=0.7正面", "其余标记为中性"],
            ["关键词", "去停用词、纯数字及无意义单字", "词云按真实词频生成"],
        ]
        self._add_table(document, ["模块", "配置", "实现要点"], method_rows, [2100, 3000, 4260])

        document.add_page_break()
        document.add_heading("4. 模型评估设计", level=1)
        if self.metrics:
            default = self.metrics["default_model"]
            custom = self.metrics["custom_model"]
            compare_rows = [
                ["默认SnowNLP", f"{default['accuracy']:.3f}", f"{default['macro_f1']:.3f}"],
                ["领域SnowNLP", f"{custom['accuracy']:.3f}", f"{custom['macro_f1']:.3f}"],
            ]
            self._add_table(document, ["模型", "Accuracy", "Macro-F1"], compare_rows, [4560, 2400, 2400])
            delta = self.metrics["macro_f1_delta"]
            if self.metrics["claim_domain_improvement"]:
                conclusion = f"领域模型Macro-F1相对默认模型提高{delta:+.3f}，达到预设的0.02声明门槛。"
            else:
                conclusion = f"领域模型Macro-F1变化为{delta:+.3f}，不足以声称性能提升；项目重点转为评估闭环和误差分析。"
            self._add_callout(document, conclusion)
            matrix = self.output_dir / "领域模型混淆矩阵.png"
            if matrix.exists():
                self._add_figure(document, matrix, "图3  领域模型在120条随机测试集上的混淆矩阵", width=5.2)
        else:
            self._add_callout(
                document,
                "已固定随机种子20260811并生成150条盲标样本：120条用于正式指标，30条用于问句、广告和阈值误差诊断。人工标签尚未填写，因此本报告不展示或虚构Accuracy、F1。",
            )
            low_questions = int(
                (
                    (self.data["情绪标签"] == -1)
                    & self.data["评论"].astype(str).map(
                        lambda text: any(marker in text for marker in QUESTION_MARKERS)
                    )
                ).sum()
            )
            self._add_body(
                document,
                f"自动风险扫描发现：当前负面预测中有{low_questions}条带有明显提问或请求特征。这不是人工真值，但提示问句语义是后续误差分析重点。",
                bold_prefix="自动风险扫描发现：",
            )

        document.add_heading("5. 关键词与内容关注点", level=1)
        top_rows = [
            [index + 1, row["关键词"], int(row["频数"])]
            for index, (_, row) in enumerate(self.keywords.head(10).iterrows())
        ]
        self._add_table(document, ["排名", "关键词", "频数"], top_rows, [1800, 4560, 3000])
        keyword_figure_no = 4 if self.metrics else 3
        cloud_figure_no = keyword_figure_no + 1
        self._add_figure(
            document,
            self.output_dir / "高频关键词_保研版.png",
            f"图{keyword_figure_no}  高频关键词排名",
        )
        self._add_figure(
            document,
            self.output_dir / "关键词词云_保研版.png",
            f"图{cloud_figure_no}  按真实频数生成的关键词词云",
            width=6.1,
        )

        document.add_heading("6. 局限性与改进方向", level=1)
        for text in (
            "数据来自单个教学视频，样本规模和课程类型有限，结果不代表普遍教学质量。",
            "SnowNLP本质上是词袋式概率模型，对提问、反问、广告和上下文依赖处理有限。",
            "训练语料只有正负两类，中性类别由阈值规则得到，需要人工评估验证。",
            "B站接口可能调整，在线采集只作为可选模块；离线数据保证分析流程可复现。",
            "后续可扩展多视频数据，并对比TF-IDF+线性模型或预训练中文模型。",
        ):
            paragraph = document.add_paragraph(style="List Bullet")
            run = paragraph.add_run(text)
            _set_run_font(run)

        document.add_heading("7. 项目结论", level=1)
        self._add_body(
            document,
            "本项目完成了从真实评论采集到领域情感分析和自动化报告的完整工程闭环。保研版重点补齐了安全配置、可复现路径、数据去重、频数词云、基线对比接口与人工评估设计，使项目能够被复核，也能诚实说明模型边界。",
        )

        output = self.output_dir / "B站教学评论情感分析项目报告_保研版.docx"
        document.save(output)
        return output

    def create_markdown_materials(self) -> None:
        if self.metrics:
            default_f1 = self.metrics["default_model"]["macro_f1"]
            custom_f1 = self.metrics["custom_model"]["macro_f1"]
            evaluation_line = (
                f"在120条随机测试集上，默认/领域模型Macro-F1分别为"
                f"{default_f1:.3f}/{custom_f1:.3f}。"
            )
        else:
            evaluation_line = "已生成150条人工评估集，当前等待完成标签，暂不声明模型准确率。"

        presentation = f"""# 3分钟项目汇报稿

老师好，我汇报的项目是“B站教学评论情感分析”。这是我独立完成的中文NLP课程项目。

项目希望解决的问题是：教学视频下有大量自然语言评论，逐条阅读效率很低，能否自动总结用户态度和关注主题。为此，我实现了从数据采集、文本处理、领域情感分析到可视化报告生成的完整流程。

数据层面，我通过B站评论接口获得了{self.summary['raw_records']}条评论，删除{self.summary['duplicates_removed']}条完全重复记录后保留{self.summary['unique_records']}条。处理阶段先做清洗，再使用Jieba、138条教学领域词典和923条停用词提取关键词。模型层面，我基于16,720条正面语料和18,976条负面语料训练SnowNLP领域模型，并采用0.3和0.7作为三分类阈值。

当前模型将评论预测为正面{self.summary['positive']}条、中性{self.summary['neutral']}条、负面{self.summary['negative']}条。新版报告不再使用原先缺乏依据的非标准满意度分数，而是保留正中负分布，并给出净情感值{self.summary['net_sentiment_points']:+.2f}个百分点。关键词方面，词云也已改为按真实词频生成。

为了验证模型，而不是只展示预测结果，我固定随机种子构建了120条随机测试样本和30条诊断样本，并实现默认SnowNLP与领域模型的Accuracy、Macro-F1和混淆矩阵对比。{evaluation_line}

项目局限是数据只来自单个视频，SnowNLP对提问、反问和广告的理解也有限。这个项目最大的收获，是完成了一条可复现的NLP工程链路，并进一步认识到模型效果必须通过独立标注和误差分析来证明。
"""
        (self.output_dir / "3分钟项目汇报稿.md").write_text(presentation, encoding="utf-8")

        claim = ""
        if self.metrics and self.metrics["claim_domain_improvement"]:
            claim = (
                f"领域模型在120条随机测试集上取得Macro-F1 "
                f"{self.metrics['custom_model']['macro_f1']:.3f}，较默认模型提高"
                f"{self.metrics['macro_f1_delta']:+.3f}。"
            )
        elif self.metrics:
            claim = "完成默认/领域模型基线对比与问句、广告、阈值样本的误差分析。"
        else:
            claim = "构建150条人工标注评估集及可复现的基线对比与误差分析流程。"

        resume = f"""# 保研简历项目表述

## 推荐项目名称

**B站教学评论情感分析与自动化报告系统｜个人课程项目**

## 简历要点

- 独立设计并实现面向B站教学视频评论的中文文本挖掘系统，完成{self.summary['raw_records']}条评论的数据采集、清洗去重、领域分词、SnowNLP情感分析及Word/PDF自动化报告生成。
- 基于35,696条正负语料训练教学领域SnowNLP模型，使用Jieba、自定义领域词典与停用词完成关键词提取；{claim}

## 面试边界

- 不把单视频结果外推为普遍教学质量结论。
- 不再使用缺乏依据的非标准满意度分数。
- 只有人工标签完成且指标支持时，才声称领域模型带来性能提升。
"""
        (self.output_dir / "保研简历项目表述.md").write_text(resume, encoding="utf-8")

    @staticmethod
    def convert_to_pdf(docx_path: Path) -> Path:
        import win32com.client

        pdf_path = docx_path.with_suffix(".pdf")
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = None
        try:
            doc = word.Documents.Open(
                str(docx_path.resolve()),
                ReadOnly=True,
                AddToRecentFiles=False,
                Visible=False,
            )
            doc.ExportAsFixedFormat(
                OutputFileName=str(pdf_path.resolve()),
                ExportFormat=17,
                OpenAfterExport=False,
                OptimizeFor=0,
                Range=0,
                Item=0,
                IncludeDocProps=True,
                KeepIRM=True,
                CreateBookmarks=1,
                DocStructureTags=True,
                BitmapMissingFonts=True,
                UseISO19005_1=False,
            )
        finally:
            if doc is not None:
                doc.Close(False)
            word.Quit()
        return pdf_path

    def start(self) -> dict[str, str]:
        self.create_charts()
        docx_path = self.create_docx()
        self.create_markdown_materials()
        result = {"docx": str(docx_path), "pdf": ""}
        try:
            result["pdf"] = str(self.convert_to_pdf(docx_path))
        except Exception as exc:
            print(f"PDF导出跳过（DOCX和图表已生成）: {exc}")
        return result


# Backward-compatible alias used by the original course project.
report = ReportGenerator


if __name__ == "__main__":
    result = ReportGenerator().start()
    print(json.dumps(result, ensure_ascii=False, indent=2))
